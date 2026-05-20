"""
docx_filler.py — DOCX 模板自动填充模块
========================================
无数据库依赖，直接传入字段值字典即可。

功能：
  1. 读取任意 DOCX 模板（支持 {{占位符}} 格式）
  2. 自动检测并转换下划线填空格式（如律师修订版框架合同）
  3. 将字段值替换进模板，保留原始格式
  4. 可选：在指定占位符位置嵌入印章图片
  5. 输出填充后的 DOCX 字节流或文件

用法：
    from contract_toolkit.docx_filler import DocxFiller

    filler = DocxFiller("template.docx")

    # 查看模板中的占位符
    print(filler.placeholders)

    # 填充字段
    filled_bytes = filler.fill({
        "乙方名称": "上海ABC贸易有限公司",
        "乙方地址": "上海市浦东新区XX路1号",
        "合同期限": "2026年1月1日至2026年12月31日",
    })

    # 保存文件
    with open("output.docx", "wb") as f:
        f.write(filled_bytes)
"""
from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Optional, Union

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.shared import Cm


class DocxFiller:
    """DOCX 模板填充器，支持 {{占位符}} 和下划线填空两种格式。"""

    PLACEHOLDER_RE = re.compile(r'\{\{([^}]+)\}\}')
    DOCX_MAGIC = b'PK\x03\x04'

    def __init__(self, template: Union[str, Path, bytes]):
        """
        Args:
            template: 模板文件路径（str/Path）或原始字节流（bytes）
        """
        if isinstance(template, (str, Path)):
            self._docx_bytes = Path(template).read_bytes()
        else:
            self._docx_bytes = template

        if not self.is_valid_docx(self._docx_bytes):
            raise ValueError("不是合法的 DOCX 文件（魔数校验失败）")

        # 如果是下划线格式，自动转换（不修改原文件）
        if self._has_underscores(self._docx_bytes):
            self._working_bytes, self._placeholders = self._convert_underscores(
                self._docx_bytes
            )
            self._converted = True
        else:
            self._working_bytes = self._docx_bytes
            self._placeholders = self._parse_placeholders(self._working_bytes)
            self._converted = False

    # ──────────────────────────────────────────────────────────────────
    # 属性
    # ──────────────────────────────────────────────────────────────────

    @property
    def placeholders(self) -> list[str]:
        """返回模板中所有占位符名称列表（去重，按首次出现顺序）。"""
        return list(self._placeholders)

    @property
    def is_converted(self) -> bool:
        """是否从下划线格式自动转换而来。"""
        return self._converted

    # ──────────────────────────────────────────────────────────────────
    # 核心填充方法
    # ──────────────────────────────────────────────────────────────────

    def fill(
        self,
        values: dict[str, str],
        seal_image: Optional[Union[str, Path, bytes]] = None,
        seal_placeholder: str = "甲方印章",
        seal_width_cm: float = 3.0,
        fallback: str = "【待填写】",
        underline_filled: bool = True,
    ) -> bytes:
        """
        用字段值填充模板，返回填充后的 DOCX 字节流。

        Args:
            values:           {占位符名称: 替换值} 字典
            seal_image:       印章图片路径或字节流（可选）
            seal_placeholder: 印章占位符名称，默认 "甲方印章"
            seal_width_cm:    印章图片宽度（厘米），默认 3.0
            fallback:         未提供的占位符替换为此文本，默认"【待填写】"
            underline_filled: 填充后的值加下划线，默认 True

        Returns:
            填充后的 DOCX 文件字节流
        """
        doc = Document(io.BytesIO(self._working_bytes))

        # 替换文本占位符
        self._replace_placeholders(doc, values, fallback, underline_filled)

        # 嵌入印章
        if seal_image is not None:
            seal_bytes = self._load_image(seal_image)
            self._embed_image(doc, f"{{{{{seal_placeholder}}}}}", seal_bytes, seal_width_cm)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def fill_to_file(
        self,
        values: dict[str, str],
        output_path: Union[str, Path],
        **kwargs,
    ) -> Path:
        """
        填充模板并保存到文件。

        Args:
            values:      字段值字典
            output_path: 输出文件路径（.docx）
            **kwargs:    透传给 fill() 的其他参数

        Returns:
            输出文件的绝对路径
        """
        filled = self.fill(values, **kwargs)
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(filled)
        return out.resolve()

    # ──────────────────────────────────────────────────────────────────
    # 静态工具方法
    # ──────────────────────────────────────────────────────────────────

    @staticmethod
    def is_valid_docx(content: bytes) -> bool:
        """校验字节流是否为合法 DOCX（ZIP 魔数）。"""
        return len(content) >= 4 and content[:4] == DocxFiller.DOCX_MAGIC

    @staticmethod
    def extract_placeholders(docx: Union[str, Path, bytes]) -> list[str]:
        """从 DOCX 文件提取所有 {{占位符}} 名称（去重）。"""
        if isinstance(docx, (str, Path)):
            data = Path(docx).read_bytes()
        else:
            data = docx
        return DocxFiller._parse_placeholders(data)

    @staticmethod
    def build_values_from_dict(
        source: dict,
        field_map: dict[str, str],
    ) -> dict[str, str]:
        """
        从任意数据字典按映射关系提取字段值。

        Args:
            source:    数据来源字典，如数据库查询结果
            field_map: {占位符名称: source字典的键名}
                       如 {"乙方名称": "company_name", "乙方地址": "address"}

        Returns:
            {占位符名称: 对应值} 字典（值为 None 的字段跳过）

        示例：
            data = {"company_name": "ABC公司", "address": "上海市..."}
            mapping = {"乙方名称": "company_name", "乙方地址": "address"}
            values = DocxFiller.build_values_from_dict(data, mapping)
            # → {"乙方名称": "ABC公司", "乙方地址": "上海市..."}
        """
        result = {}
        for placeholder, key in field_map.items():
            val = source.get(key)
            if val is not None and str(val).strip():
                result[placeholder] = str(val)
        return result

    # ──────────────────────────────────────────────────────────────────
    # 内部方法
    # ──────────────────────────────────────────────────────────────────

    @staticmethod
    def _parse_placeholders(docx_bytes: bytes) -> list[str]:
        try:
            doc = Document(io.BytesIO(docx_bytes))
        except Exception:
            return []

        found = []
        re_ = DocxFiller.PLACEHOLDER_RE

        def _scan(paragraphs):
            for para in paragraphs:
                text = "".join(r.text for r in para.runs)
                for m in re_.finditer(text):
                    found.append(m.group(1))

        _scan(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _scan(cell.paragraphs)
        for section in doc.sections:
            for part in [section.header, section.footer,
                         section.first_page_header, section.first_page_footer,
                         section.even_page_header, section.even_page_footer]:
                if part:
                    _scan(part.paragraphs)
        return list(dict.fromkeys(found))

    @staticmethod
    def _has_underscores(docx_bytes: bytes) -> bool:
        try:
            doc = Document(io.BytesIO(docx_bytes))
        except Exception:
            return False
        all_text = "".join(p.text for p in doc.paragraphs)
        return bool(re.search(r'_{2,}', all_text))

    @staticmethod
    def _convert_underscores(
        docx_bytes: bytes,
        placeholder_order: Optional[list[str]] = None,
    ) -> tuple[bytes, list[str]]:
        """将下划线填空转换为 {{...}} 占位符格式。"""
        DEFAULT_ORDER = [
            "甲方法定代表人", "甲方联系方式", "甲方信用代码",
            "乙方名称", "乙方地址", "乙方法定代表人", "乙方联系电话", "统一社会信用代码",
            "合同额度",
            "合同期限_起_年", "合同期限_起_月", "合同期限_起_日",
            "合同期限_止_年", "合同期限_止_月", "合同期限_止_日",
            "包装指导书编号",
            "甲方签字", "签订日期_年", "签订日期_月", "签订日期_日",
            "乙方印章", "乙方签字",
            "乙方签字日期_年", "乙方签字日期_月", "乙方签字日期_日",
        ]
        order = placeholder_order or DEFAULT_ORDER
        ph_iter = iter(order)
        used: list[str] = []

        doc = Document(io.BytesIO(docx_bytes))

        def _replace(para) -> bool:
            full = "".join(r.text for r in para.runs)
            if not re.search(r'_{2,}', full):
                return False

            def sub(m):
                try:
                    ph = next(ph_iter)
                    used.append(ph)
                    return "{{" + ph + "}}"
                except StopIteration:
                    return m.group(0)

            new = re.sub(r'_{2,}', sub, full)
            if new != full and para.runs:
                for r in para.runs:
                    r.text = ""
                para.runs[0].text = new
                return True
            return False

        for p in doc.paragraphs:
            _replace(p)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _replace(p)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue(), list(dict.fromkeys(used))

    @staticmethod
    def _replace_placeholders(
        doc: Document, values: dict[str, str], fallback: str, underline_filled: bool = True
    ) -> None:
        re_ = DocxFiller.PLACEHOLDER_RE

        def _fill(paragraphs):
            for para in paragraphs:
                runs = para.runs
                if not runs:
                    continue
                full = "".join(r.text for r in runs)
                if not re_.search(full):
                    continue

                # 清空所有 run 文本
                for r in runs:
                    r.text = ""

                if not underline_filled:
                    # 原始行为：不加下划线，直接替换
                    new = re_.sub(lambda m: values.get(m.group(1), fallback), full)
                    runs[0].text = new
                    return

                # 按占位符拆分为片段：[(文本, 是否为填充值), ...]
                segments: list[tuple[str, bool]] = []
                last_end = 0
                for m in re_.finditer(full):
                    if m.start() > last_end:
                        segments.append((full[last_end:m.start()], False))
                    value = values.get(m.group(1), fallback)
                    segments.append((value, True))
                    last_end = m.end()
                if last_end < len(full):
                    segments.append((full[last_end:], False))

                # 写回 run：第一段复用 runs[0]，后续用 para.add_run 新建
                first_written = False
                for text, is_value in segments:
                    if not text:
                        continue
                    if not first_written:
                        run = runs[0]
                        run.text = text
                        first_written = True
                    else:
                        run = para.add_run(text)

                    # 填充值加下划线，非填充部分明确不加（避免继承）
                    run.underline = True if is_value else None

        _fill(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _fill(cell.paragraphs)

    @staticmethod
    def _load_image(source: Union[str, Path, bytes]) -> bytes:
        if isinstance(source, bytes):
            return source
        return Path(source).read_bytes()

    @staticmethod
    def _embed_image(
        doc: Document, placeholder_text: str, image_bytes: bytes, width_cm: float
    ) -> bool:
        buf = io.BytesIO(image_bytes)

        def _try(paragraphs) -> bool:
            for para in paragraphs:
                full = "".join(r.text for r in para.runs)
                if placeholder_text not in full:
                    continue
                for r in para.runs:
                    r.text = ""
                run = para.add_run()
                run.add_picture(buf, width=Cm(width_cm))
                return True
            return False

        if _try(doc.paragraphs):
            return True
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if _try(cell.paragraphs):
                        return True
        return False
